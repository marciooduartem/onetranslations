import os
import json
import base64
import requests
import uuid
import threading
from pathlib import Path
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from pdf2docx import Converter
import fitz

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = Path("uploads")
OUTPUT_FOLDER = Path("outputs")
UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)

# ── ESTADO DOS JOBS (em memória) ──────────────────────────────────
JOBS = {}  # job_id → {status, current, total, message, output_path, error}

# ── PROMPTS ───────────────────────────────────────────────────────
PROMPT_SIMPLES_TABELAS = """Extraia APENAS as tabelas desta página e retorne SOMENTE JSON válido.
{"tabelas":[{"linhas":[["col1","col2"],["dado1","dado2"]]}]}

REGRAS CRÍTICAS:
- Leia a tabela SEMPRE de cima para baixo, da esquerda para a direita
- Primeira linha de cada tabela é o cabeçalho
- Preserve TODAS as células, mesmo se vazias (use "" para célula vazia)
- Se há células mescladas no cabeçalho, repita o valor mesclado em cada célula
- Preserve exatamente o texto com acentos, símbolos e pontuação
- NÃO inclua texto fora de tabelas
- Se não houver tabelas na página, retorne {"tabelas":[]}"""

PROMPT_JURAMENTADO = """Você é especialista em formatação de traduções juramentadas brasileiras.
Analise a imagem e retorne SOMENTE JSON válido:
{"idioma":"pt"|"en"|"es"|"de"|"fr"|"it"|"zh","blocos":[
  {"tipo":"linha","texto":"..."},
  {"tipo":"vazia"},
  {"tipo":"tabela","titulo":"Table 1...","linhas":[["col1","col2"],["dado1","dado2"]]}
]}

LEITURA DE TABELAS — CRÍTICO:
- Leia SEMPRE de cima para baixo, da esquerda para a direita
- Primeira linha do array "linhas" é o cabeçalho
- Cada linha seguinte é uma linha de dados
- Preserve TODAS as células (use "" se vazia)
- Células mescladas no cabeçalho: repita o valor em cada célula que ela cobre

SUBSTITUIÇÕES VISUAIS OBRIGATÓRIAS:
- Logotipo legível → bloco "linha" com "[Consta logotipo]" + blocos seguintes com texto do logo
- Brasão/armas → "[Consta brasão de (país/estado)]"
- Selo oficial → "[Consta selo de (país/estado/instituição)]"
- Foto de pessoa → "[Consta fotografia]"
- Imagem única → "[Consta imagem]"
- Múltiplas imagens → "[Constam imagens]"
- Figura técnica com texto → "[Consta figura]" + linhas com texto
- Desenho técnico → "[Consta desenho técnico]" ou "[Constam desenhos técnicos]"
- Gráfico → "[Consta gráfico]"
- Diagrama → "[Consta diagrama]"
- QR Code → "[Consta código QR]"
- Assinatura → "[Consta assinatura]" ou "[Constam assinaturas]"
- Carimbo legível → "[Consta carimbo]" + linhas com texto do carimbo
- Carimbo ilegível → "[Consta carimbo ilegível]"
- Texto ilegível → "[Ilegível]"
- Palavras riscadas → NÃO incluir

REGRAS DE FORMATAÇÃO:
- Itens numerados (1., 2., 3.) em UMA LINHA SÓ: "1. Trademark: FERRARI"
- Subitens (16.1., 16.2.) também em uma linha
- Preserve acentos, maiúsculas, símbolos especiais"""

PROMPT_CETRA_EXTRA = """
MODO CETRA ATIVADO:
- Documentos bilíngues: manter APENAS o texto em inglês
- Exceções: cabeçalho e nomes de cargos/referências institucionais
- Palavras riscadas: NUNCA incluir
- Rodapé repetitivo: incluir APENAS na última página, precedido de "[consta nota de rodapé]"
- Imagens com título "Drawing No. X": incluir título + citação
- Numeração: corrigir pontos faltantes (1.11 → 1.1.1)"""

# ── EXTRAÇÃO VIA CLAUDE API ───────────────────────────────────────
def extract_page_claude(image_bytes, api_key, system_prompt):
    b64 = base64.b64encode(image_bytes).decode()
    resp = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={"Content-Type": "application/json", "x-api-key": api_key, "anthropic-version": "2023-06-01"},
        json={
            "model": "claude-opus-4-5",
            "max_tokens": 4000,
            "system": system_prompt,
            "messages": [{"role": "user", "content": [
                {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": b64}},
                {"type": "text", "text": "Retorne SOMENTE o JSON."}
            ]}]
        },
        timeout=120
    )
    data = resp.json()
    if not resp.ok:
        raise Exception(data.get("error", {}).get("message", "Erro na API"))
    raw = data["content"][0]["text"].strip().replace("```json","").replace("```","").strip()
    return json.loads(raw)

# ── BUILDER JURAMENTADO ───────────────────────────────────────────
def build_docx_juramentado(all_pages, output_path):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(12)

    def add_text(texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)
        if texto.strip() == "":
            p.paragraph_format.space_after = Pt(8)
        else:
            run = p.add_run(texto)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            stripped = texto.strip()
            if stripped.isupper() and len(stripped) < 80 and not stripped.startswith("["):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_table_real(titulo, linhas):
        if titulo:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(titulo)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
        if not linhas: return
        num_cols = max(len(r) for r in linhas)
        if num_cols == 0: return
        table = doc.add_table(rows=len(linhas), cols=num_cols)
        table.style = "Table Grid"
        for ri, linha in enumerate(linhas):
            for ci in range(num_cols):
                cell = table.cell(ri, ci)
                cell.text = ""
                texto = linha[ci] if ci < len(linha) else ""
                if ri == 0:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "D9D9D9")
                    tcPr.append(shd)
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(texto)
                run.font.name = "Arial"
                run.font.size = Pt(10)
                if ri == 0:
                    run.bold = True
        doc.add_paragraph()

    first = True
    for page_data in all_pages:
        if not first:
            doc.add_page_break()
        first = False
        blocos = page_data.get("blocos", [])
        # Compatibilidade formato antigo
        if not blocos and page_data.get("linhas"):
            for l in page_data["linhas"]:
                blocos.append({"tipo": "vazia" if l.strip() == "" else "linha", "texto": l})
        for bloco in blocos:
            tipo = bloco.get("tipo", "linha")
            if tipo == "vazia":
                add_text("")
            elif tipo == "tabela":
                add_table_real(bloco.get("titulo", ""), bloco.get("linhas", []))
            else:
                add_text(bloco.get("texto", ""))
    doc.save(str(output_path))

# ── BUILDER SIMPLES (pdf2docx + remoção de carimbos/assinaturas) ─
def build_docx_simples(pdf_path, output_path):
    """
    Usa pdf2docx para preservar tudo (incluindo imagens originais).
    Carimbos/assinaturas serão removidos manualmente pelo usuário ou via IA opcional.
    """
    cv = Converter(str(pdf_path))
    cv.convert(str(output_path))
    cv.close()

# ── PROCESSAMENTO EM BACKGROUND ───────────────────────────────────
def process_juramentado_async(job_id, pdf_path, output_path, api_key, cetra):
    try:
        doc = fitz.open(str(pdf_path))
        total = len(doc)
        JOBS[job_id]["total"] = total
        all_pages = []
        system = PROMPT_JURAMENTADO + (PROMPT_CETRA_EXTRA if cetra else "")

        for page_num in range(total):
            JOBS[job_id]["current"] = page_num + 1
            JOBS[job_id]["message"] = f"Processando página {page_num + 1} de {total}..."
            page = doc[page_num]
            mat = fitz.Matrix(1.5, 1.5)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("jpeg")
            try:
                data = extract_page_claude(img_bytes, api_key, system)
                all_pages.append(data)
            except Exception as e:
                all_pages.append({"blocos": [{"tipo": "linha", "texto": f"[Erro na página {page_num+1}: {str(e)[:100]}]"}]})
        doc.close()

        JOBS[job_id]["message"] = "Gerando arquivo Word..."
        build_docx_juramentado(all_pages, output_path)
        JOBS[job_id]["status"] = "done"
        JOBS[job_id]["message"] = "Concluído!"
    except Exception as e:
        JOBS[job_id]["status"] = "error"
        JOBS[job_id]["error"] = str(e)
    finally:
        if pdf_path.exists():
            try: pdf_path.unlink()
            except: pass

def process_simples_async(job_id, pdf_path, output_path):
    try:
        JOBS[job_id]["message"] = "Convertendo PDF preservando layout, imagens e tabelas..."
        JOBS[job_id]["total"] = 1
        JOBS[job_id]["current"] = 0
        build_docx_simples(pdf_path, output_path)
        JOBS[job_id]["current"] = 1
        JOBS[job_id]["status"] = "done"
        JOBS[job_id]["message"] = "Concluído!"
    except Exception as e:
        JOBS[job_id]["status"] = "error"
        JOBS[job_id]["error"] = str(e)
    finally:
        if pdf_path.exists():
            try: pdf_path.unlink()
            except: pass

# ── ROTAS ─────────────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/convert", methods=["POST"])
def convert():
    if "file" not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400
    file = request.files["file"]
    mode = request.form.get("mode", "simples")
    cetra = request.form.get("cetra", "false") == "true"
    api_key = request.form.get("api_key", "")

    job_id = str(uuid.uuid4())
    pdf_path = UPLOAD_FOLDER / f"{job_id}.pdf"
    output_path = OUTPUT_FOLDER / f"{job_id}.docx"
    file.save(str(pdf_path))

    original_name = Path(file.filename).stem
    JOBS[job_id] = {
        "status": "processing",
        "current": 0,
        "total": 0,
        "message": "Iniciando...",
        "output_path": str(output_path),
        "filename": f"{original_name}_convertido.docx",
        "error": None
    }

    if mode == "juramentado":
        if not api_key:
            del JOBS[job_id]
            return jsonify({"error": "A tradução juramentada requer a chave de API."}), 400
        thread = threading.Thread(target=process_juramentado_async, args=(job_id, pdf_path, output_path, api_key, cetra))
    else:
        thread = threading.Thread(target=process_simples_async, args=(job_id, pdf_path, output_path))

    thread.daemon = True
    thread.start()
    return jsonify({"job_id": job_id})

@app.route("/status/<job_id>")
def status(job_id):
    job = JOBS.get(job_id)
    if not job:
        return jsonify({"error": "Job não encontrado"}), 404
    return jsonify({
        "status": job["status"],
        "current": job["current"],
        "total": job["total"],
        "message": job["message"],
        "error": job.get("error")
    })

@app.route("/download/<job_id>")
def download(job_id):
    job = JOBS.get(job_id)
    if not job or job["status"] != "done":
        return jsonify({"error": "Arquivo não disponível"}), 404
    return send_file(
        job["output_path"],
        as_attachment=True,
        download_name=job["filename"],
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
